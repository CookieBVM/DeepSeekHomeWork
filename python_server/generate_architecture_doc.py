# -*- coding: utf-8 -*-
"""
生成项目架构文档 - .docx 格式
使用 python-docx 库生成专业的 Word 文档
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def set_chinese_font(run):
    """设置中文字体"""
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def add_heading(doc, text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        set_chinese_font(run)
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 128)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 0)
        elif level == 3:
            run.font.size = Pt(12)
            run.font.bold = True
    return heading


def add_paragraph(doc, text, bold=False, size=11, indent=True):
    """添加段落"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run)
    run.font.size = Pt(size)
    run.font.bold = bold
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    return para


def add_normal_text(doc, text):
    """添加普通文本（无缩进）"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_chinese_font(run)
    run.font.size = Pt(11)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(3)
    return para


def add_list_item(doc, text, bullet=False, level=0):
    """添加列表项"""
    para = doc.add_paragraph(style='List Bullet' if bullet else 'List Number')
    run = para.add_run(text)
    set_chinese_font(run)
    run.font.size = Pt(11)
    para.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(3)
    return para


def add_code_block(doc, code_text):
    """添加代码块"""
    para = doc.add_paragraph()
    run = para.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 128)
    para.paragraph_format.left_indent = Cm(0.74)
    para.paragraph_format.line_spacing = 1.2
    para.paragraph_format.space_after = Pt(6)
    return para


def add_separator(doc):
    """添加分隔线"""
    para = doc.add_paragraph()
    run = para.add_run("_" * 80)
    run.font.color.rgb = RGBColor(192, 192, 192)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)


def create_document():
    """创建 Word 文档"""
    
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    title = doc.add_heading('Unity + Python 具身多模态交互项目', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_chinese_font(run)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 128)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Python 端架构与技术原理文档')
    set_chinese_font(run)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 128, 128)
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f'文档生成日期: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    set_chinese_font(run)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_page_break()
    
    add_heading(doc, '目录', level=1)
    add_list_item(doc, '第一章 项目概述', bullet=False)
    add_list_item(doc, '第二章 整体架构设计', bullet=False)
    add_list_item(doc, '第三章 Python 与 Unity 通信原理', bullet=False)
    add_list_item(doc, '第四章 核心模块详解', bullet=False)
    add_list_item(doc, '第五章 数据处理原理', bullet=False)
    add_list_item(doc, '第六章 离线缓存机制', bullet=False)
    add_list_item(doc, '第七章 家长端报告生成', bullet=False)
    add_list_item(doc, '第八章 通信协议规范', bullet=False)
    add_list_item(doc, '附录', bullet=False)
    
    doc.add_page_break()
    
    add_heading(doc, '第一章 项目概述', level=1)
    
    add_paragraph(doc, '本项目是一个基于 Unity + Python 架构的具身多模态交互系统，专为自闭症儿童设计，提供低压力的社交练习环境。项目包含两个核心场景：买菜对话和合作涂色。')
    
    add_heading(doc, '项目目标', level=2)
    add_list_item(doc, '帮助自闭症儿童通过游戏化方式练习社交技能', bullet=True)
    add_list_item(doc, '提供语音交互和触控交互两种交互方式', bullet=True)
    add_list_item(doc, '实时数据分析，提供正反馈鼓励机制', bullet=True)
    add_list_item(doc, '家长端数据同步，便于了解孩子的学习进度', bullet=True)
    
    add_heading(doc, '技术特点', level=2)
    add_list_item(doc, 'Python 3.11 兼容，PEP 8 编码规范', bullet=True)
    add_list_item(doc, '标准库优先，无外部依赖', bullet=True)
    add_list_item(doc, '支持两种通信模式：标准输入输出 / 文件监听', bullet=True)
    add_list_item(doc, '基于 Levenshtein 距离的模糊匹配算法', bullet=True)
    add_list_item(doc, '完整的离线缓存机制', bullet=True)
    
    doc.add_page_break()
    
    add_heading(doc, '第二章 整体架构设计', level=1)
    
    add_heading(doc, '2.1 系统层次结构', level=2)
    add_paragraph(doc, '系统采用分层架构设计，Unity 负责前端展示和用户交互，Python 负责后端逻辑和数据处理。')
    
    add_normal_text(doc, '                    +---------------------+')
    add_normal_text(doc, '                    |    Unity (C#)      |')
    add_normal_text(doc, '                    |  UI渲染 / 动画 /    |')
    add_normal_text(doc, '                    |  触控交互           |')
    add_normal_text(doc, '                    +---------+-----------+')
    add_normal_text(doc, '                              |')
    add_normal_text(doc, '                    +---------v-----------+')
    add_normal_text(doc, '                    |  通信桥接层         |')
    add_normal_text(doc, '                    |  PythonBridge.cs   |')
    add_normal_text(doc, '                    +---------+-----------+')
    add_normal_text(doc, '                              |')
    add_normal_text(doc, '           +------------------+------------------+')
    add_normal_text(doc, '           |                                     |')
    add_normal_text(doc, '  +--------v-------+                    +--------v-------+')
    add_normal_text(doc, '  | stdin/stdout   |                    |  文件监听      |')
    add_normal_text(doc, '  |  进程间通信     |                    |  (JSON文件)    |')
    add_normal_text(doc, '  +--------+-------+                    +--------+-------+')
    add_normal_text(doc, '           |                                     |')
    add_normal_text(doc, '           +------------------+------------------+')
    add_normal_text(doc, '                              |')
    add_normal_text(doc, '                    +---------v-----------+')
    add_normal_text(doc, '                    |  Python 服务端      |')
    add_normal_text(doc, '                    |  main.py (入口)    |')
    add_normal_text(doc, '                    +---------+-----------+')
    add_normal_text(doc, '                              |')
    add_normal_text(doc, '       +----------------------+----------------------+')
    add_normal_text(doc, '       |                      |                      |')
    add_normal_text(doc, '  +----v----+          +----v----+          +----v----+')
    add_normal_text(doc, '  |语音处理  |          |数据记录  |          |报告生成  |')
    add_normal_text(doc, '  |模块      |          |分析模块  |          |模块      |')
    add_normal_text(doc, '  +---------+          +---------+          +---------+')
    
    add_heading(doc, '2.2 Python 端模块组成', level=2)
    add_paragraph(doc, 'Python 端由 6 个核心文件组成：')
    
    add_heading(doc, '1. config.py - 全局配置文件', level=3)
    add_list_item(doc, '语音识别配置（模糊匹配阈值、置信度阈值）', bullet=True)
    add_list_item(doc, '场景意图库定义（买菜场景、涂色场景）', bullet=True)
    add_list_item(doc, '数据记录路径配置', bullet=True)
    add_list_item(doc, 'Unity 通信配置', bullet=True)
    
    add_heading(doc, '2. speech_processor.py - 语音交互模块', level=3)
    add_list_item(doc, 'SpeechProcessor 类：语音转文本和意图分析', bullet=True)
    add_list_item(doc, 'transcribe_audio()：模拟语音识别，支持模糊匹配', bullet=True)
    add_list_item(doc, 'parse_intent()：基于关键词匹配的意图识别', bullet=True)
    add_list_item(doc, '基于 Levenshtein 距离的模糊匹配算法', bullet=True)
    
    add_heading(doc, '3. data_recorder.py - 数据记录与分析模块', level=3)
    add_list_item(doc, 'InteractionEvent 数据类：交互事件模型', bullet=True)
    add_list_item(doc, 'SessionStats 数据类：会话统计模型', bullet=True)
    add_list_item(doc, 'DataRecorder 类：事件记录和实时统计', bullet=True)
    add_list_item(doc, '离线缓存机制（JSON/Pickle 格式）', bullet=True)
    
    add_heading(doc, '4. report_generator.py - 家长端报告模块', level=3)
    add_list_item(doc, 'ParentReport 数据类：报告数据模型', bullet=True)
    add_list_item(doc, 'ReportGenerator 类：报告生成和导出', bullet=True)
    add_list_item(doc, 'generate_summary()：自然语言总结生成', bullet=True)
    add_list_item(doc, 'export_for_parent()：JSON/TXT 格式导出', bullet=True)
    
    add_heading(doc, '5. main.py - 主入口与通信接口', level=3)
    add_list_item(doc, 'UnityBridge 类：与 Unity 通信的桥接器', bullet=True)
    add_list_item(doc, '支持标准输入输出和文件监听两种模式', bullet=True)
    add_list_item(doc, '消息分发与处理', bullet=True)
    
    add_heading(doc, '6. test_modules.py - 测试脚本', level=3)
    add_list_item(doc, '模块功能验证', bullet=True)
    add_list_item(doc, '集成测试', bullet=True)
    
    doc.add_page_break()
    
    add_heading(doc, '第三章 Python 与 Unity 通信原理', level=1)
    
    add_heading(doc, '3.1 通信架构概述', level=2)
    add_paragraph(doc, 'Unity 与 Python 采用跨进程通信方式，提供两种通信模式。')
    
    add_heading(doc, '模式一：标准输入输出 (stdin/stdout)', level=3)
    add_list_item(doc, 'Unity 启动 Python 子进程', bullet=True)
    add_list_item(doc, '通过重定向 stdin/stdout 进行 JSON 数据交换', bullet=True)
    add_list_item(doc, '优点：实时性好、无需文件系统', bullet=True)
    add_list_item(doc, '适用：常规运行环境', bullet=True)
    
    add_heading(doc, '模式二：文件监听 (File Watch)', level=3)
    add_list_item(doc, 'Unity 写入请求到 JSON 文件', bullet=True)
    add_list_item(doc, 'Python 定时轮询读取并处理', bullet=True)
    add_list_item(doc, 'Python 写入响应到另一个 JSON 文件', bullet=True)
    add_list_item(doc, 'Unity 读取响应文件', bullet=True)
    add_list_item(doc, '优点：调试方便、可独立运行', bullet=True)
    add_list_item(doc, '适用：调试模式、无网络环境', bullet=True)
    
    add_heading(doc, '3.2 标准输入输出模式原理', level=2)
    add_paragraph(doc, '工作流程：')
    
    add_heading(doc, '1. Unity 端 (PythonBridge.cs)', level=3)
    add_list_item(doc, '使用 Process.Start() 启动 python main.py', bullet=True)
    add_list_item(doc, '设置 RedirectStandardInput = true', bullet=True)
    add_list_item(doc, '设置 RedirectStandardOutput = true', bullet=True)
    add_list_item(doc, '设置 RedirectStandardError = true', bullet=True)
    add_list_item(doc, '设置 UseShellExecute = false', bullet=True)
    
    add_heading(doc, '2. 发送请求', level=3)
    add_list_item(doc, 'Unity 将请求对象序列化为 JSON', bullet=True)
    add_list_item(doc, '通过 process.StandardInput.WriteLine() 发送', bullet=True)
    add_list_item(doc, '消息以换行符 \"\\n\" 分隔', bullet=True)
    
    add_heading(doc, '3. Python 处理', level=3)
    add_list_item(doc, 'main.py 中 sys.stdin.readline() 循环读取', bullet=True)
    add_list_item(doc, 'json.loads() 解析请求', bullet=True)
    add_list_item(doc, '调用对应的 handler 处理', bullet=True)
    
    add_heading(doc, '4. 返回响应', level=3)
    add_list_item(doc, 'Python 将响应字典序列化为 JSON', bullet=True)
    add_list_item(doc, 'print(json.dumps(response), flush=True) 输出', bullet=True)
    add_list_item(doc, 'flush=True 确保立即发送', bullet=True)
    
    add_heading(doc, '5. Unity 接收', level=3)
    add_list_item(doc, '使用 Task.Run() 异步读取 StandardOutput', bullet=True)
    add_list_item(doc, '解析 JSON 响应', bullet=True)
    add_list_item(doc, '通过 UnityMainThreadDispatcher 回调到主线程', bullet=True)
    
    add_heading(doc, '3.3 文件监听模式原理', level=2)
    
    add_heading(doc, '1. 目录结构', level=3)
    add_normal_text(doc, '  python_server/data/communication/')
    add_normal_text(doc, '    - unity_to_python.json (Unity 写入，Python 读取)')
    add_normal_text(doc, '    - python_to_unity.json (Python 写入，Unity 读取)')
    
    add_heading(doc, '2. 发送请求', level=3)
    add_list_item(doc, 'Unity 调用 File.WriteAllText() 写入 JSON', bullet=True)
    add_list_item(doc, 'Python 每 0.5 秒检查文件修改时间', bullet=True)
    add_list_item(doc, '发现文件变化则读取并处理', bullet=True)
    
    add_heading(doc, '3. 返回响应', level=3)
    add_list_item(doc, 'Python 处理完成后写入输出文件', bullet=True)
    add_list_item(doc, '然后删除输入文件（可选）', bullet=True)
    add_list_item(doc, 'Unity 读取输出文件', bullet=True)
    
    add_heading(doc, '3.4 异步输出读取机制', level=2)
    add_paragraph(doc, '关键实现细节：')
    
    add_list_item(doc, '异步读取：使用 Task.Run(() => ReadOutputAsync())', bullet=True)
    add_list_item(doc, '缓冲机制：StringBuilder 累积字符直到遇到换行符', bullet=True)
    add_list_item(doc, '行处理：逐行解析，支持流数据', bullet=True)
    add_list_item(doc, '跨线程安全：UnityMainThreadDispatcher.Enqueue() 确保 UI 更新在主线程', bullet=True)
    
    add_paragraph(doc, '为什么需要异步读取？')
    add_list_item(doc, '阻塞读取会导致 Unity 主线程冻结', bullet=True)
    add_list_item(doc, 'Python 可能持续输出日志', bullet=True)
    add_list_item(doc, '需要实时响应', bullet=True)
    
    doc.add_page_break()
    
    add_heading(doc, '第四章 核心模块详解', level=1)
    
    add_heading(doc, '4.1 配置模块 (config.py)', level=2)
    add_paragraph(doc, '配置模块定义了项目的所有常量和设置：')
    
    add_heading(doc, '语音识别配置：', level=3)
    add_list_item(doc, 'fuzzy_match_threshold: 3 (Levenshtein 距离阈值)', bullet=True)
    add_list_item(doc, 'confidence_threshold: 0.6 (意图识别置信度阈值)', bullet=True)
    
    add_heading(doc, '意图库设计：', level=3)
    add_list_item(doc, '买菜场景：buy_item, ask_price, greet, thanks, bye', bullet=True)
    add_list_item(doc, '涂色场景：color_object, select_object, finish', bullet=True)
    add_list_item(doc, '每个意图包含：keywords (触发词), entities (可提取实体)', bullet=True)
    
    add_heading(doc, '实体定义示例：', level=3)
    add_code_block(doc, '"苹果": ["苹果", "pingguo", "pinguo", "ping guo", "apple"]')
    add_list_item(doc, '支持中文、拼音、空格、英文多种表达方式', bullet=True)
    add_list_item(doc, '模糊匹配允许发音不清的情况', bullet=True)
    
    add_heading(doc, '4.2 语音处理模块 (speech_processor.py)', level=2)
    
    add_paragraph(doc, '核心类：SpeechProcessor')
    
    add_heading(doc, 'transcribe_audio(audio_path) 方法：', level=3)
    add_list_item(doc, '模拟语音转文本功能', bullet=True)
    add_list_item(doc, '根据文件名映射到预设文本', bullet=True)
    add_list_item(doc, '调用 _apply_noise_resistance() 进行清洗', bullet=True)
    
    add_heading(doc, '_apply_noise_resistance(text) 方法：', level=3)
    add_list_item(doc, '转小写、去首尾空格', bullet=True)
    add_list_item(doc, '移除标点符号', bullet=True)
    add_list_item(doc, '过滤口语填充词（嗯嗯、那个、这个、啊、呃、嗯）', bullet=True)
    add_list_item(doc, '模拟儿童发音不清晰的处理', bullet=True)
    
    add_heading(doc, 'parse_intent(text, context) 方法算法流程：', level=3)
    add_list_item(doc, '1. 检查文本是否为空', bullet=True)
    add_list_item(doc, '2. 根据 context 选择意图库', bullet=True)
    add_list_item(doc, '3. 遍历所有意图，计算匹配得分', bullet=True)
    add_list_item(doc, '   a. _calculate_keyword_score(): 关键词匹配得分', bullet=True)
    add_list_item(doc, '   b. _extract_entities(): 提取实体', bullet=True)
    add_list_item(doc, '   c. _calculate_confidence(): 综合置信度', bullet=True)
    add_list_item(doc, '4. 选择置信度最高的意图', bullet=True)
    add_list_item(doc, '5. 低于阈值则返回 unknown', bullet=True)
    
    add_heading(doc, '4.3 数据记录模块 (data_recorder.py)', level=2)
    
    add_heading(doc, 'InteractionEvent (交互事件) 数据模型：', level=3)
    add_list_item(doc, 'event_id: 事件唯一标识', bullet=True)
    add_list_item(doc, 'event_type: 事件类型', bullet=True)
    add_list_item(doc, 'data: 事件数据字典', bullet=True)
    add_list_item(doc, 'timestamp: 时间戳', bullet=True)
    add_list_item(doc, 'session_id: 会话ID', bullet=True)
    add_list_item(doc, 'context: 场景上下文', bullet=True)
    add_list_item(doc, 'correctness: 正确性判断', bullet=True)
    
    add_heading(doc, 'SessionStats (会话统计) 数据模型：', level=3)
    add_list_item(doc, 'session_id: 会话ID', bullet=True)
    add_list_item(doc, 'start_time/end_time: 时间范围', bullet=True)
    add_list_item(doc, 'total_events: 总事件数', bullet=True)
    add_list_item(doc, 'correct_events: 正确事件数', bullet=True)
    add_list_item(doc, 'wrong_events: 错误事件数', bullet=True)
    add_list_item(doc, 'consecutive_correct: 连续正确次数', bullet=True)
    add_list_item(doc, 'max_consecutive_correct: 最高连续正确', bullet=True)
    add_list_item(doc, 'total_duration: 总时长', bullet=True)
    add_list_item(doc, 'scene_stats: 各场景统计', bullet=True)
    add_list_item(doc, 'accuracy (属性): 正确率计算', bullet=True)
    
    add_heading(doc, 'log_event() 方法流程：', level=3)
    add_list_item(doc, '1. 生成事件ID和时间戳', bullet=True)
    add_list_item(doc, '2. _determine_correctness(): 判断正确性', bullet=True)
    add_list_item(doc, '3. 创建 InteractionEvent 对象', bullet=True)
    add_list_item(doc, '4. 添加到事件列表', bullet=True)
    add_list_item(doc, '5. _update_stats(): 更新统计', bullet=True)
    add_list_item(doc, '6. _save_to_cache(): 保存到缓存', bullet=True)
    add_list_item(doc, '7. _append_to_log(): 追加到日志', bullet=True)
    
    add_heading(doc, '_determine_correctness() 规则：', level=3)
    add_list_item(doc, '如果 data 包含 is_correct，直接使用', bullet=True)
    add_list_item(doc, 'intent_detected: confidence >= 0.7 为正确', bullet=True)
    add_list_item(doc, 'user_action: action == expected_action', bullet=True)
    add_list_item(doc, 'color_selection: selected_color == expected_color', bullet=True)
    add_list_item(doc, 'item_selection: selected_item == expected_item', bullet=True)
    
    add_heading(doc, '正反馈触发机制：', level=3)
    add_list_item(doc, 'streak_threshold: 3 (连续正确次数阈值)', bullet=True)
    add_list_item(doc, '当 consecutive_correct >= streak_threshold 时', bullet=True)
    add_list_item(doc, 'should_trigger_feedback = True', bullet=True)
    add_list_item(doc, 'Unity 根据此标志播放庆祝动画、给予贴纸等奖励', bullet=True)
    
    doc.add_page_break()
    
    add_heading(doc, '第五章 数据处理原理', level=1)
    
    add_heading(doc, '5.1 模糊匹配算法原理', level=2)
    add_paragraph(doc, '项目使用 Levenshtein 距离（编辑距离）实现模糊匹配。')
    
    add_heading(doc, '什么是 Levenshtein 距离？', level=3)
    add_list_item(doc, '定义：将一个字符串转换为另一个字符串所需的最少操作数', bullet=True)
    add_list_item(doc, '操作包括：插入、删除、替换一个字符', bullet=True)
    
    add_paragraph(doc, '示例：')
    add_list_item(doc, '"pingguo" -> "pinguo": 替换一个字符，距离 = 1', bullet=True)
    add_list_item(doc, '"hongse" -> "红色": 完全不同，距离 = 6', bullet=True)
    
    add_heading(doc, '算法实现（动态规划）：', level=3)
    add_list_item(doc, '1. 创建 (n+1) x (m+1) 的矩阵', bullet=True)
    add_list_item(doc, '2. 初始化第一行和第一列为索引值', bullet=True)
    add_list_item(doc, '3. 逐行逐列填充：', bullet=True)
    add_list_item(doc, '   - 如果字符相等，取左上角值', bullet=True)
    add_list_item(doc, '   - 否则，取 min(插入, 删除, 替换) + 1', bullet=True)
    add_list_item(doc, '4. 右下角值即为编辑距离', bullet=True)
    
    add_heading(doc, '_fuzzy_match() 多层匹配策略：', level=3)
    add_list_item(doc, '1. 精确子串匹配：pattern in text', bullet=True)
    add_list_item(doc, '2. 忽略大小写：pattern_lower in text_lower', bullet=True)
    add_list_item(doc, '3. 按单词分割匹配：对每个单词计算距离', bullet=True)
    add_list_item(doc, '4. 滑动窗口匹配：对长模式进行子串距离计算', bullet=True)
    
    add_heading(doc, '5.2 意图识别算法', level=2)
    add_paragraph(doc, '基于规则的关键词匹配算法：')
    
    add_heading(doc, '关键词得分计算：', level=3)
    add_list_item(doc, '遍历意图的所有关键词', bullet=True)
    add_list_item(doc, '使用模糊匹配判断是否命中', bullet=True)
    add_list_item(doc, '得分 = 命中数 / 关键词总数', bullet=True)
    add_list_item(doc, '范围：0.0 ~ 1.0', bullet=True)
    
    add_heading(doc, '实体提取：', level=3)
    add_list_item(doc, '遍历每个实体类型（如 items, colors）', bullet=True)
    add_list_item(doc, '对每个实体值（如 苹果、香蕉）', bullet=True)
    add_list_item(doc, '检查其所有别名（如 "苹果", "pingguo", "apple"）', bullet=True)
    add_list_item(doc, '模糊匹配成功则提取该实体', bullet=True)
    
    add_heading(doc, '置信度计算：', level=3)
    add_code_block(doc, 'confidence = keyword_score * 0.5 + entity_score * 0.5')
    
    add_paragraph(doc, '实体得分：')
    add_list_item(doc, '每个实体类型贡献最高 1.0 分', bullet=True)
    add_list_item(doc, '每个实体值贡献 0.3 分（上限 1.0）', bullet=True)
    add_list_item(doc, '多实体类型取平均', bullet=True)
    
    add_heading(doc, '最终决策：', level=3)
    add_list_item(doc, '选择所有意图中置信度最高的', bullet=True)
    add_list_item(doc, '如果最高置信度 < 0.6，标记为 unknown', bullet=True)
    
    add_heading(doc, '5.3 实时统计原理', level=2)
    add_paragraph(doc, '统计数据在每次事件记录时实时更新：')
    
    add_heading(doc, '_update_stats() 方法：', level=3)
    add_list_item(doc, '1. total_events += 1', bullet=True)
    add_list_item(doc, '2. 根据 correctness 更新：', bullet=True)
    add_list_item(doc, '   - CORRECT: correct_events++, consecutive_correct++', bullet=True)
    add_list_item(doc, '   - WRONG: wrong_events++, consecutive_correct = 0', bullet=True)
    add_list_item(doc, '3. 更新 max_consecutive_correct', bullet=True)
    add_list_item(doc, '4. 更新 scene_stats（按场景分类统计）', bullet=True)
    add_list_item(doc, '5. 更新 total_duration', bullet=True)
    
    add_heading(doc, 'get_current_stats() 返回的实时数据：', level=3)
    add_list_item(doc, 'accuracy: 正确率（动态计算）', bullet=True)
    add_list_item(doc, 'total_duration: 当前时长', bullet=True)
    add_list_item(doc, 'consecutive_correct: 当前连续正确数', bullet=True)
    add_list_item(doc, 'should_trigger_feedback: 是否触发正反馈', bullet=True)
    add_list_item(doc, 'scene_stats: 各场景独立统计', bullet=True)
    
    add_heading(doc, 'Unity 如何使用这些数据？', level=3)
    add_list_item(doc, '每记录一个事件就调用 get_stats', bullet=True)
    add_list_item(doc, '检查 should_trigger_feedback', bullet=True)
    add_list_item(doc, '如果为 True：播放庆祝动画、显示贴纸、播放音乐', bullet=True)
    add_list_item(doc, '显示实时正确率给孩子（可选）', bullet=True)
    
    doc.add_page_break()
    
    add_heading(doc, '第六章 离线缓存机制', level=1)
    
    add_heading(doc, '6.1 缓存目录结构', level=2)
    add_normal_text(doc, '  python_server/data/')
    add_normal_text(doc, '    - cache/              # 会话缓存')
    add_normal_text(doc, '      - session_20240115_143022.json')
    add_normal_text(doc, '      - session_20240115_150010.pkl')
    add_normal_text(doc, '    - logs/               # 事件日志')
    add_normal_text(doc, '      - events_20240115.log')
    add_normal_text(doc, '    - reports/            # 家长端报告')
    add_normal_text(doc, '      - report_session_xxx.json')
    add_normal_text(doc, '      - report_session_xxx.txt')
    add_normal_text(doc, '    - communication/      # 文件监听模式')
    add_normal_text(doc, '      - unity_to_python.json')
    add_normal_text(doc, '      - python_to_unity.json')
    
    add_heading(doc, '6.2 缓存数据结构', level=2)
    add_paragraph(doc, '缓存文件包含：')
    add_code_block(doc, '''{
    "session_id": "session_20240115_143022",
    "events": [ 所有事件列表 ],
    "stats": { 会话统计数据 },
    "last_updated": 时间戳
}''')
    
    add_heading(doc, '6.3 缓存流程', level=2)
    
    add_heading(doc, '保存流程 (_save_to_cache)：', level=3)
    add_list_item(doc, '1. 构造缓存数据字典', bullet=True)
    add_list_item(doc, '2. 序列化：', bullet=True)
    add_list_item(doc, '   - JSON 格式：json.dump(ensure_ascii=False)', bullet=True)
    add_list_item(doc, '   - Pickle 格式：pickle.dump()', bullet=True)
    add_list_item(doc, '3. 写入文件', bullet=True)
    add_list_item(doc, '4. 每次 log_event 后立即保存', bullet=True)
    
    add_heading(doc, '加载流程 (_load_cached_data)：', level=3)
    add_list_item(doc, '1. 检查缓存文件是否存在', bullet=True)
    add_list_item(doc, '2. 读取并反序列化', bullet=True)
    add_list_item(doc, '3. _restore_from_cache() 恢复数据', bullet=True)
    add_list_item(doc, '4. 初始化时自动调用', bullet=True)
    
    add_heading(doc, '6.4 会话恢复', level=2)
    add_paragraph(doc, '场景：网络中断后重连')
    add_list_item(doc, 'Python 进程可能被重启', bullet=True)
    add_list_item(doc, 'DataRecorder 初始化时加载缓存', bullet=True)
    add_list_item(doc, '恢复所有事件和统计数据', bullet=True)
    add_list_item(doc, '可以继续之前的会话', bullet=True)
    
    add_heading(doc, 'load_session() 方法：', level=3)
    add_list_item(doc, '读取指定 session_id 的缓存', bullet=True)
    add_list_item(doc, '恢复 DataRecorder 状态', bullet=True)
    add_list_item(doc, '支持查看历史数据', bullet=True)
    add_list_item(doc, '支持批量报告生成', bullet=True)
    
    add_heading(doc, '6.5 日志系统', level=2)
    
    add_heading(doc, '日志文件格式：', level=3)
    add_list_item(doc, '每行一个 JSON 对象', bullet=True)
    add_list_item(doc, '按天分割（events_YYYYMMDD.log）', bullet=True)
    add_list_item(doc, '包含完整事件信息', bullet=True)
    
    add_heading(doc, '日志内容：', level=3)
    add_list_item(doc, 'timestamp: ISO 格式时间', bullet=True)
    add_list_item(doc, 'event_id, event_type, session_id', bullet=True)
    add_list_item(doc, 'context, correctness', bullet=True)
    add_list_item(doc, 'data: 完整事件数据', bullet=True)
    
    add_heading(doc, '用途：', level=3)
    add_list_item(doc, '问题排查', bullet=True)
    add_list_item(doc, '数据回放', bullet=True)
    add_list_item(doc, '分析工具导入', bullet=True)
    
    doc.add_page_break()
    
    add_heading(doc, '第七章 家长端报告生成', level=1)
    
    add_heading(doc, '7.1 报告数据模型', level=2)
    add_paragraph(doc, 'ParentReport 数据类包含：')
    add_list_item(doc, 'session_id: 会话ID', bullet=True)
    add_list_item(doc, 'report_date: 报告生成时间', bullet=True)
    add_list_item(doc, 'summary_text: 自然语言总结', bullet=True)
    add_list_item(doc, 'accuracy: 正确率', bullet=True)
    add_list_item(doc, 'total_duration: 总时长', bullet=True)
    add_list_item(doc, 'correct_events/wrong_events: 正确/错误数', bullet=True)
    add_list_item(doc, 'max_consecutive_correct: 最高连续正确', bullet=True)
    add_list_item(doc, 'scene_performance: 各场景表现', bullet=True)
    add_list_item(doc, 'strengths: 优势领域', bullet=True)
    add_list_item(doc, 'areas_to_improve: 改进方向', bullet=True)
    
    add_heading(doc, '7.2 自然语言总结生成', level=2)
    add_paragraph(doc, 'generate_summary() 策略：')
    
    add_heading(doc, '1. 开场白', level=3)
    add_paragraph(doc, '"宝宝今天在游戏中表现很棒！"')
    
    add_heading(doc, '2. 正确率评价（分级）', level=3)
    add_list_item(doc, '>= 80%: "正确率达到了 X%，完成了 Y 次正确互动"', bullet=True)
    add_list_item(doc, '>= 60%: "正确率为 X%，完成了 Y 次正确互动，继续加油"', bullet=True)
    add_list_item(doc, '>= 40%: "正确率为 X%，完成了 Y 次正确互动，多多练习"', bullet=True)
    add_list_item(doc, '< 40%: "今天完成了 Y 次互动，一起加油"', bullet=True)
    
    add_heading(doc, '3. 连续正确表扬', level=3)
    add_list_item(doc, '>= 5: "特别棒！最高连续正确 X 次！"', bullet=True)
    add_list_item(doc, '>= 3: "很好！最高连续正确 X 次。"', bullet=True)
    
    add_heading(doc, '4. 时长评价', level=3)
    add_list_item(doc, '>= 10 分钟: "非常专注"', bullet=True)
    add_list_item(doc, '>= 5 分钟: "表现不错"', bullet=True)
    add_list_item(doc, '< 5 分钟: "下次可以多玩一会儿"', bullet=True)
    
    add_heading(doc, '5. 场景细分评价', level=3)
    add_list_item(doc, '每个场景独立评价', bullet=True)
    add_list_item(doc, '显示 correct/total 和百分比', bullet=True)
    
    add_heading(doc, '7.3 优势与改进分析', level=2)
    
    add_heading(doc, '_identify_strengths() 规则：', level=3)
    add_list_item(doc, '整体正确率 >= 70%: "整体表现优秀"', bullet=True)
    add_list_item(doc, '最高连续正确 >= 5: "能够持续专注"', bullet=True)
    add_list_item(doc, '某场景正确率 >= 80% 且互动 >= 3: "在XX中表现出色"', bullet=True)
    add_list_item(doc, '总时长 >= 600 秒: "能够长时间专注"', bullet=True)
    add_list_item(doc, '默认: "正在积极参与"', bullet=True)
    
    add_heading(doc, '_identify_areas_to_improve() 规则：', level=3)
    add_list_item(doc, '整体正确率 < 50%: "可以多加练习"', bullet=True)
    add_list_item(doc, '某场景正确率 < 40% 且互动 >= 2: "建议多加练习XX场景"', bullet=True)
    add_list_item(doc, '总时长 < 180 秒: "可以尝试延长游戏时间"', bullet=True)
    
    add_heading(doc, '7.4 报告导出', level=2)
    
    add_heading(doc, 'JSON 格式 (_export_json)：', level=3)
    add_list_item(doc, '完整数据序列化', bullet=True)
    add_list_item(doc, '便于程序解析', bullet=True)
    add_list_item(doc, '适合家长端 App 读取', bullet=True)
    
    add_heading(doc, 'TXT 格式 (_export_txt)：', level=3)
    add_list_item(doc, '人类可读格式', bullet=True)
    add_list_item(doc, '包含分隔线和标题', bullet=True)
    add_list_item(doc, '便于家长直接阅读', bullet=True)
    add_list_item(doc, '结尾带有鼓励话语', bullet=True)
    
    add_heading(doc, 'simulate_send_to_parent()：', level=3)
    add_list_item(doc, '模拟网络发送', bullet=True)
    add_list_item(doc, '可扩展为真实 API 调用', bullet=True)
    
    doc.add_page_break()
    
    add_heading(doc, '第八章 通信协议规范', level=1)
    
    add_heading(doc, '8.1 请求格式', level=2)
    add_paragraph(doc, 'Unity 发送的 JSON 请求：')
    add_code_block(doc, '''{
    "action": "操作名称",
    "session_id": "可选，会话ID",
    ... 其他参数 ...
}''')
    
    add_heading(doc, '8.2 响应格式', level=2)
    add_paragraph(doc, 'Python 返回的 JSON 响应：')
    add_code_block(doc, '''{
    "success": true/false,
    "action": "对应的操作名称",
    "data": { ... 返回数据 ... },
    "error": "错误消息（失败时）",
    "timestamp": 时间戳
}''')
    
    add_heading(doc, '8.3 支持的操作 (Action)', level=2)
    
    add_heading(doc, '1. ping - 心跳检测', level=3)
    add_paragraph(doc, '请求:', indent=False)
    add_code_block(doc, '{"action": "ping"}')
    add_paragraph(doc, '响应:', indent=False)
    add_code_block(doc, '{ "success": true, "data": { "status": "alive", "session_id": "..." } }')
    
    add_heading(doc, '2. transcribe_audio - 语音转文本', level=3)
    add_paragraph(doc, '请求:', indent=False)
    add_code_block(doc, '{"action": "transcribe_audio", "audio_path": "xxx.wav"}')
    
    add_heading(doc, '3. parse_intent - 意图分析', level=3)
    add_paragraph(doc, '请求:', indent=False)
    add_code_block(doc, '''{"action": "parse_intent", "text": "买苹果", "context": "buying_vegetables"}''')
    add_paragraph(doc, '副作用：自动记录 intent_detected 事件')
    
    add_heading(doc, '4. log_event - 记录事件', level=3)
    add_paragraph(doc, '请求:', indent=False)
    add_code_block(doc, '''{"action": "log_event", 
 "event_type": "color_selection",
 "data": {"selected_color": "红色", "expected_color": "红色"},
 "context": "coloring"}''')
    
    add_heading(doc, '5. get_stats - 获取统计', level=3)
    add_paragraph(doc, '请求:', indent=False)
    add_code_block(doc, '{"action": "get_stats"}')
    add_paragraph(doc, '返回数据包含：', indent=False)
    add_list_item(doc, 'session_id, accuracy, total_duration', bullet=True)
    add_list_item(doc, 'total_events, correct_events, wrong_events', bullet=True)
    add_list_item(doc, 'consecutive_correct, max_consecutive_correct', bullet=True)
    add_list_item(doc, 'should_trigger_feedback, scene_stats', bullet=True)
    
    add_heading(doc, '6. generate_report - 生成报告', level=3)
    add_paragraph(doc, '请求:', indent=False)
    add_code_block(doc, '{"action": "generate_report"}')
    
    add_heading(doc, '7. export_report - 导出报告', level=3)
    add_paragraph(doc, '请求:', indent=False)
    add_code_block(doc, '{"action": "export_report", "format": "json"/"txt"}')
    
    add_heading(doc, '8. end_session - 结束会话', level=3)
    add_paragraph(doc, '请求:', indent=False)
    add_code_block(doc, '{"action": "end_session"}')
    add_paragraph(doc, '返回：final_stats + report', indent=False)
    
    add_heading(doc, '9. list_sessions - 列出会话', level=3)
    add_paragraph(doc, '返回所有可用的会话ID列表', indent=False)
    
    add_heading(doc, '10. load_session - 加载会话', level=3)
    add_paragraph(doc, '加载指定历史会话数据', indent=False)
    
    add_heading(doc, '8.4 典型交互流程', level=2)
    add_paragraph(doc, '买菜场景完整流程：')
    
    add_list_item(doc, '1. Unity: 启动 Python 服务 -> Process.Start("python main.py")', bullet=True)
    add_list_item(doc, '2. Unity: 发送心跳 ping', bullet=True)
    add_list_item(doc, '3. Unity: 记录场景开始 log_event(event_type="scene_start")', bullet=True)
    add_list_item(doc, '4. 孩子说："买苹果"（Unity 录音）', bullet=True)
    add_list_item(doc, '5. Unity: 语音转文本 transcribe_audio', bullet=True)
    add_list_item(doc, '6. Unity: 意图分析 parse_intent("买苹果", "buying_vegetables")', bullet=True)
    add_list_item(doc, '7. Unity: 根据意图执行动作（数字人拿苹果）', bullet=True)
    add_list_item(doc, '8. Unity: 记录正确事件 log_event(item_selection)', bullet=True)
    add_list_item(doc, '9. Unity: 检测到 should_trigger_feedback = true -> 播放庆祝动画', bullet=True)
    add_list_item(doc, '10. 游戏结束', bullet=True)
    add_list_item(doc, '11. Unity: 结束会话并获取报告 end_session', bullet=True)
    add_list_item(doc, '12. Unity: 导出报告给家长 export_report', bullet=True)
    
    doc.add_page_break()
    
    add_heading(doc, '附录', level=1)
    
    add_heading(doc, 'A. 文件位置参考', level=2)
    
    add_heading(doc, 'Python 文件：', level=3)
    add_list_item(doc, r'd:\DeepSeek\python_server\main.py', bullet=True)
    add_list_item(doc, r'd:\DeepSeek\python_server\config.py', bullet=True)
    add_list_item(doc, r'd:\DeepSeek\python_server\speech_processor.py', bullet=True)
    add_list_item(doc, r'd:\DeepSeek\python_server\data_recorder.py', bullet=True)
    add_list_item(doc, r'd:\DeepSeek\python_server\report_generator.py', bullet=True)
    add_list_item(doc, r'd:\DeepSeek\python_server\test_modules.py', bullet=True)
    
    add_heading(doc, 'Unity C# 文件：', level=3)
    add_list_item(doc, r'd:\DeepSeek\Assets\DeepSeek\Scripts\PythonBridge.cs', bullet=True)
    add_list_item(doc, r'd:\DeepSeek\Assets\DeepSeek\Scripts\PythonBridgeTester.cs', bullet=True)
    
    add_heading(doc, 'B. 运行测试', level=2)
    
    add_heading(doc, 'Python 端测试：', level=3)
    add_code_block(doc, '''cd python_server
python test_modules.py''')
    
    add_heading(doc, 'Unity 端测试：', level=3)
    add_list_item(doc, '1. 在 Unity 中添加 PythonBridgeTester 组件', bullet=True)
    add_list_item(doc, '2. 配置 Python 路径和脚本路径', bullet=True)
    add_list_item(doc, '3. 运行场景，点击测试按钮', bullet=True)
    
    add_heading(doc, 'C. 技术栈', level=2)
    
    add_heading(doc, 'Python 端：', level=3)
    add_list_item(doc, 'Python 3.11+', bullet=True)
    add_list_item(doc, '标准库：json, os, sys, time, pickle, dataclasses, typing, re, datetime', bullet=True)
    add_list_item(doc, '无第三方依赖', bullet=True)
    
    add_heading(doc, 'Unity 端：', level=3)
    add_list_item(doc, 'Unity 2021+', bullet=True)
    add_list_item(doc, 'Newtonsoft.Json (用于 JSON 序列化)', bullet=True)
    add_list_item(doc, 'C# 8.0+', bullet=True)
    
    add_heading(doc, 'D. 扩展方向', level=2)
    
    add_list_item(doc, '1. 接入真实 ASR 服务：百度语音、讯飞、Whisper 等，修改 transcribe_audio() 方法', bullet=True)
    add_list_item(doc, '2. 接入真实 NLP 服务：大语言模型 API，修改 parse_intent() 方法', bullet=True)
    add_list_item(doc, '3. 家长端真实发送：实现 HTTP API 调用，修改 simulate_send_to_parent()', bullet=True)
    add_list_item(doc, '4. 更多场景：扩展 config.py 中的意图库，添加更多意图和实体', bullet=True)
    add_list_item(doc, '5. 数据分析可视化：使用 matplotlib 或 Plotly 生成图表报告', bullet=True)
    
    return doc


def main():
    """主函数"""
    print("正在生成 Word 文档...")
    
    doc = create_document()
    
    output_path = os.path.join(
        os.path.dirname(__file__),
        "Python_端架构与技术原理文档.docx"
    )
    
    doc.save(output_path)
    
    print(f"\n文档已生成: {output_path}")
    print(f"\n使用说明：")
    print(f"  1. 双击 .docx 文件可用 Microsoft Word 2007+ 打开")
    print(f"  2. 如需 .doc 格式，在 Word 中选择'另存为' -> 'Word 97-2003 文档 (*.doc)'")
    print(f"  3. 文档包含完整的样式和格式")


if __name__ == "__main__":
    main()
